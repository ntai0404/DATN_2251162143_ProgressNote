import os
import time
import sys
import glob
from playwright.sync_api import sync_playwright
from pathlib import Path
from docx import Document
from dotenv import load_dotenv

# Đảm bảo in Tiếng Việt chuẩn và Flush ngay lập tức
if sys.stdout.encoding != 'utf-8':
    import io
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

def log_print(msg):
    print(f"[{time.strftime('%H:%M:%S')}] {msg}", flush=True)

load_dotenv()

BASE_DIR = Path(__file__).parent.parent
DOWNLOAD_DIR = BASE_DIR / "data_raw" / "tvpl"
DOWNLOAD_DIR.mkdir(parents=True, exist_ok=True)

TVPL_USER = "P2"
TVPL_PASS = "dhtl123456"

# Danh sách 7 mục tiêu chuẩn với ID chính xác 100%
targets = [
    ("01_Luat_GDDH_2012", "https://thuvienphapluat.vn/van-ban/Giao-duc/Luat-giao-duc-dai-hoc-2012-141398.aspx"),
    ("02_Luat_Suadoi_2018", "https://thuvienphapluat.vn/van-ban/Giao-duc/Luat-sua-doi-Luat-Giao-duc-dai-hoc-2018-371747.aspx"),
    ("03_Thongtu_08_2021", "https://thuvienphapluat.vn/van-ban/Giao-duc/Thong-tu-08-2021-TT-BGDDT-quy-che-dao-tao-trinh-do-dai-hoc-468368.aspx"),
    ("04_Nghidinh_99_2019", "https://thuvienphapluat.vn/van-ban/Giao-duc/Nghi-dinh-99-2019-ND-CP-huong-dan-Luat-Giao-duc-dai-hoc-432906.aspx"),
    ("05_Thongtu_10_2016", "https://thuvienphapluat.vn/van-ban/Giao-duc/Thong-tu-10-2016-TT-BGDDT-Quy-che-cong-tac-sinh-vien-dai-hoc-he-chinh-quy-307513.aspx"),
    ("06_Luat_Vienchuc_2010", "https://thuvienphapluat.vn/van-ban/Bo-may-hanh-chinh/Luat-Vien-chuc-2010-107730.aspx"),
    ("07_Luat_Giaoduc_2019", "https://thuvienphapluat.vn/van-ban/Giao-duc/Luat-Giao-duc-2019-333172.aspx"),
]

def harvest_core(p, nickname, url):
    """Xử lý từng file trong một session riêng biệt để tránh kẹt Cloudflare"""
    browser = p.chromium.launch(headless=True)
    context = browser.new_context(accept_downloads=True, user_agent="Mozilla/5.0 (Windows NT 10.0; Win64; x64)")
    page = context.new_page()
    
    try:
        log_print(f"🚀 Processing: {nickname}...")
        # 1. Login
        page.goto("https://thuvienphapluat.vn/page/login.aspx", timeout=45000)
        page.fill("#UserName", TVPL_USER)
        page.fill("#Password", TVPL_PASS)
        page.click("#Button1")
        page.wait_for_timeout(3000)
        
        # 2. Navigate to document
        page.goto(url, timeout=60000)
        page.wait_for_timeout(7000)

        # 3. Strategy: Scrape DOM directly (Most stable for all 7)
        # We target the actual content div, even if it's currently hidden behind a tab.
        content_text = page.evaluate('''() => {
            // First, trigger content tab if possible to be safe
            const tabs = Array.from(document.querySelectorAll('li, a, span'));
            const contentTab = tabs.find(e => e.innerText.trim().includes("Nội dung") || e.innerText.trim().includes("gốc"));
            if(contentTab) contentTab.click();
            
            // Wait slightly for DOM to settle (simulated)
            const contentEl = document.querySelector('#ctl00_Content_ThongTinVB_divNoiDung, .content-vb, body');
            if(contentEl) {
                // Remove clutter before returning
                const garbage = contentEl.querySelectorAll('header, footer, nav, .sidebar, .ads, .popup');
                garbage.forEach(g => g.remove());
                return contentEl.innerText;
            }
            return "";
        }''')
        
        if len(content_text) > 1000:
            doc = Document()
            doc.add_heading(nickname.replace("_", " "), 0)
            for line in content_text.split('\n'):
                if line.strip(): doc.add_paragraph(line.strip())
            
            out_path = DOWNLOAD_DIR / f"{nickname}.doc"
            doc.save(str(out_path))
            log_print(f"   +++ ✅ SUCCESS: {nickname}.doc ({len(content_text)} chars)")
            return True
        else:
            log_print(f"   --- ❌ FAILED: Content too short for {nickname}")
            return False

    except Exception as e:
        log_print(f"   --- ❌ Error: {str(e)[:100]}")
        return False
    finally:
        browser.close()

def main():
    log_print(f"=== TVPL SURGICAL HARVESTER v2.0 (All-In-One) ===")
    with sync_playwright() as p:
        success_count = 0
        for nick, url in targets:
            if harvest_core(p, nick, url):
                success_count += 1
            time.sleep(3) # Anti-ban delay

    log_print("\n" + "="*60)
    log_print(f"FINAL RESULT: {success_count} / 7 documents collected.")
    log_print("="*60)

if __name__ == "__main__":
    main()
